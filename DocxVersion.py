import PyE_tools as pye
import math
import decimal
import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

data = []
dataToShow = []
exercise = []
documentName = "problem.txt"
resultDocumentName = "result.docx"

# Lectura de datos
try:
    print(f'Abriendo: {documentName}...')
    info = open(documentName, "r", -1, "utf-8").readlines()
except:
    print("No se ha podido abrir el archivo problem.txt :(")
    quit()


# Lecturas de documento problem.txt
isGivingData = False
isDescribing = False

print(f'Leyendo información de {documentName}...')
for i in info:
    if i.strip() == "DE:":
        isDescribing = True
        isGivingData = False
        continue
    elif i.strip() == "DA:":
        isDescribing = False
        isGivingData = True
        continue
    
    if(isGivingData):
        try:
            data.append(decimal.Decimal(i.strip()))
            dataToShow.append((i.strip()))
        except:
            print(f'No hemos podido convertir el siguiente dato: "{i.strip()}", lo omitiremos por ahora :)')
    elif(isDescribing):
        exercise.append(i.strip())

# Escribiendo el documento
document = docx.Document()

# Fuente
font = document.styles["Normal"].font
font.name = "Arial"

# Si no hay datos
if len(data) == 0:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.bold = True
    run.add_text("¿Cómo se supone que se calcule algo que si no hay datos?")
    print(f"Guardando en: {resultDocumentName}...")
    document.save(resultDocumentName)
    quit()
    
# Escritura del problema
if len(exercise) != 0:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.bold = True
    run.add_text("Problema: ")
    run = paragraph.add_run()
    run.add_text("\n".join(exercise))


# Escritura de los primero datos
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.bold = True
run.add_text("Datos: ")
run = paragraph.add_run()
run.add_text(", ".join(dataToShow) + ".")

# Escritura de los datos ordenados
print("Haciendo cálculos poderosos...")
dataSorted = pye.bubbleSort(data)
dataToShow = []
k = pye.klassesNumber(data)
r = pye.dataRange(dataSorted)
a = pye.dataAmplitudeByList(dataSorted, True)
uv = pye.variationUnit(dataSorted)
dataTable = pye.calculateFrecuencyByDataList(dataSorted,k,a,uv)

for i in dataSorted:
    dataToShow.append(str(i))

paragraph = document.add_paragraph()
run = paragraph.add_run()
run.bold = True
run.add_text("Datos Ordenados: ")
run = paragraph.add_run()
run.add_text(", ".join(dataToShow) + ".")

# Número de clases
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.bold = True
run.add_text("Número de Clases: ")
run = paragraph.add_run()
res = (1 + 3.322 * math.log10(len(data)))
res = res if type(res) is int else "%.3f" % res
run.add_text(
    f"1 + 3.322 x log({len(dataSorted)}) = {res} ≈ {k}"
)

# Rango
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.bold = True
run.add_text("Rango: ")
run = paragraph.add_run()
run.add_text(
    f"{str(dataSorted[len(dataSorted)-1])} − {str(dataSorted[0])} = {str(dataSorted[len(dataSorted)-1] - dataSorted[0])} ≈ {r}"
)

# Amplitud
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.bold = True
run.add_text("Amplitud: ")
run = paragraph.add_run()
res = r/k
res = res if type(res) is int else "%.3f" % res
run.add_text(f"{r} ÷ {k} = {res} ≈ {a}")

# Unidad de variacion
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.bold = True
run.add_text("Unidad de variación: ")
run = paragraph.add_run()
run.add_text(str(uv))


# Tabla
print("Creando tabla...")
header = [
    "Clase",
    "Límite Inf.",
    "Límite Sup.",
    "Frec.",
    "Marca d Clases",
    "Lim Inf Exac",
    "Lim Sup Exac",
]
table = document.add_table(rows=k + 1, cols=len(header))
table.autofit = True

for i in range (0,len(header)):
    table.cell(0,i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.cell(0, i).text = header[i]
    
for i in range(0, len(dataTable)):
    for k in range(0, len(dataTable[i])):
        table.cell(i+1,k).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.cell(i + 1, k).text = str(dataTable[i][k])

print(f"Guardando en: {resultDocumentName}...")
document.save(resultDocumentName)